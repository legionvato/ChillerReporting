import json
import os
import io
from datetime import date
import streamlit as st
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook

# Config
st.set_page_config(page_title="Trane Chiller Maintenance Report", layout="wide")
CHECKLIST_PATH = os.path.join("docs", "checklists", "trane_chiller_v1.json")
FOOTER_TEXT = "Treimax Georgia Maintenance / Service Reporting Tool"
STATUS_OPTIONS = ["", "OK", "Not OK", "N/A"]

# Helpers
def safe_text(x: str) -> str:
    return (x or "").strip()

def load_checklist(path: str) -> tuple[dict, str]:
    if not os.path.exists(path):
        return {"name": "Checklist missing", "version": "0", "sections": []}, f"Checklist file not found: {path}"
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data, f"Loaded checklist from: {path}"
    except Exception as e:
        return {"name": "Checklist read error", "version": "0", "sections": []}, f"Failed to read checklist: {e}"

def is_removed_item(item: dict) -> bool:
    item_id = (item.get("id") or "").lower()
    label = (item.get("label") or "").lower()
    if item_id == "safety_loto":
        return True
    if "loto" in label or "ppe" in label:
        return True
    return False

def normalize_sections(checklist: dict) -> list[dict]:
    sections = checklist.get("sections") or []
    out = []
    for s in sections:
        title = s.get("title") or s.get("name") or "Section"
        items = s.get("items") or []
        items = [it for it in items if not is_removed_item(it)]
        out.append({"title": title, "items": items})
    return out

def compute_summary(report: dict) -> dict:
    counts = {"OK": 0, "Not OK": 0, "N/A": 0, "": 0}
    not_ok_items = []
    final_status_not_ok = False
    id_to_label = {}
    for sec in report["sections"]:
        for it in sec["items"]:
            id_to_label[it["id"]] = it["label"]
    for item_id, v in (report.get("results") or {}).items():
        status = (v or {}).get("status", "")
        if status not in counts:
            counts[""] += 1
        else:
            counts[status] += 1
        if status == "Not OK":
            label = id_to_label.get(item_id, item_id)
            not_ok_items.append(label)
    for item_id, label in id_to_label.items():
        status = (report.get("results") or {}).get(item_id, {}).get("status", "")
        lid = (item_id or "").lower()
        ll = (label or "").lower()
        if status == "Not OK" and (
            "final" in lid
            or "returned to normal operation" in ll
            or "final status" in ll
            or "unit returned" in ll
        ):
            final_status_not_ok = True
            break
    overall = "OK"
    if counts["Not OK"] > 0:
        overall = "ATTENTION"
    if final_status_not_ok:
        overall = "CRITICAL"
    return {"counts": counts, "not_ok_items": not_ok_items, "overall": overall}

def validate_report(report: dict) -> list[str]:
    errors = []
    id_to_label = {}
    for sec in report["sections"]:
        for it in sec["items"]:
            id_to_label[it["id"]] = it["label"]
    for item_id, v in (report.get("results") or {}).items():
        status = (v or {}).get("status", "")
        notes = safe_text((v or {}).get("notes", ""))
        if status == "Not OK" and not notes:
            errors.append(f'Notes required for "Not OK": {id_to_label.get(item_id, item_id)}')
    return errors

# PDF Export
class PDFWriter:
    def __init__(self, c: canvas.Canvas, title: str):
        self.c = c
        self.title = title
        self.w, self.h = A4
        self.margin_l = self.margin_r = 16 * mm
        self.margin_t = self.margin_b = 16 * mm
        self.header_h = 14 * mm
        self.footer_h = 10 * mm
        self.content_top = self.h - self.margin_t - self.header_h
        self.content_bottom = self.margin_b + self.footer_h + 2 * mm
        self.x0 = self.margin_l
        self.x1 = self.w - self.margin_r
        self.max_w = self.x1 - self.x0
        self.y = self.content_top
        self._draw_header()

    def _draw_header(self):
        c = self.c
        c.saveState()
        bar_y = self.h - self.margin_t - self.header_h
        c.setFillColor(colors.whitesmoke)
        c.rect(self.x0, bar_y, self.max_w, self.header_h, stroke=0, fill=1)
        c.setFillColor(colors.black)
        c.setFont("Helvetica-Bold", 13)
        c.drawString(self.x0 + 4 * mm, bar_y + 4.5 * mm, self.title)
        c.restoreState()
        self.y = self.content_top - 2 * mm

    def _draw_footer(self):
        c = self.c
        c.saveState()
        y = self.margin_b + self.footer_h
        c.setStrokeColor(colors.lightgrey)
        c.setLineWidth(0.5)
        c.line(self.x0, y, self.x1, y)
        c.setFillColor(colors.grey)
        c.setFont("Helvetica", 8)
        c.drawCentredString((self.x0 + self.x1) / 2, self.margin_b + 3.5 * mm, FOOTER_TEXT)
        c.drawRightString(self.x1, self.margin_b + 3.5 * mm, f"Page {c.getPageNumber()}")
        c.restoreState()

    def _new_page(self):
        self._draw_footer()
        self.c.showPage()
        self._draw_header()

    def _ensure_space(self, needed_h: float):
        if self.y - needed_h < self.content_bottom:
            self._new_page()

    def _text(self, txt: str, size=10, bold=False, color=colors.black, leading=12):
        self._ensure_space(leading)
        self.c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        self.c.setFillColor(color)
        self.c.drawString(self.x0, self.y, txt)
        self.y -= leading * 0.9

    def _text_wrapped(self, txt: str, size=10, bold=False, color=colors.black, leading=12, indent=0):
        txt = txt or ""
        c = self.c
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.setFillColor(color)
        max_w = self.max_w - indent
        words = txt.split()
        if not words:
            self._ensure_space(leading)
            self.y -= leading * 0.9
            return
        line = []
        while words:
            line.append(words.pop(0))
            if words:
                w = c.stringWidth(" ".join(line + [words[0]]), c._fontname, size)
                if w > max_w:
                    self._ensure_space(leading)
                    c.drawString(self.x0 + indent, self.y, " ".join(line))
                    self.y -= leading * 0.9
                    line = []
        if line:
            self._ensure_space(leading)
            c.drawString(self.x0 + indent, self.y, " ".join(line))
            self.y -= leading * 0.9

    def _section_bar(self, title: str):
        bar_h = 8 * mm
        self._ensure_space(bar_h + 4 * mm)
        y = self.y - bar_h + 2 * mm
        self.c.setFillColor(colors.HexColor("#F2F2F2"))
        self.c.rect(self.x0, y, self.max_w, bar_h, stroke=0, fill=1)
        self.c.setFillColor(colors.black)
        self.c.setFont("Helvetica-Bold", 11)
        self.c.drawString(self.x0 + 3 * mm, y + 2.5 * mm, title)
        self.y = y - 3 * mm

    def _badge(self, status: str, x_right: float, y_baseline: float):
        status = status or "—"
        c = self.c
        if status == "OK":
            fill, stroke, textc = colors.HexColor("#EAF4EA"), colors.HexColor("#7A9A7A"), colors.HexColor("#2F5F2F")
        elif status == "Not OK":
            fill, stroke, textc = colors.HexColor("#F7EAEA"), colors.HexColor("#A05A5A"), colors.HexColor("#7A1F1F")
        elif status == "N/A":
            fill, stroke, textc = colors.HexColor("#F0F0F0"), colors.HexColor("#9A9A9A"), colors.HexColor("#5A5A5A")
        else:
            fill, stroke, textc = colors.white, colors.HexColor("#C0C0C0"), colors.HexColor("#555555")
        c.saveState()
        c.setFont("Helvetica-Bold", 9)
        pad_x = 3 * mm
        pad_y = 1.6 * mm
        text_w = c.stringWidth(status, "Helvetica-Bold", 9)
        box_w = text_w + 2 * pad_x
        box_h = 6.5 * mm
        x = x_right - box_w
        y = y_baseline - 1.8 * mm
        c.setFillColor(fill)
        c.setStrokeColor(stroke)
        c.setLineWidth(0.8)
        c.roundRect(x, y, box_w, box_h, radius=2.2 * mm, stroke=1, fill=1)
        c.setFillColor(textc)
        c.drawString(x + pad_x, y + pad_y, status)
        c.restoreState()

    def _item_line(self, label: str, status: str):
        right_reserve = 30 * mm
        c = self.c
        c.setFont("Helvetica", 10)
        max_w = self.max_w - right_reserve
        words = (label or "").split()
        lines = []
        line = []
        while words:
            line.append(words.pop(0))
            if words:
                w = c.stringWidth(" ".join(line + [words[0]]), "Helvetica", 10)
                if w > max_w:
                    lines.append(" ".join(line))
                    line = [words.pop(0)] if words else []
        if line:
            lines.append(" ".join(line))
        leading = 12
        total_h = leading * len(lines) + 3 * mm
        self._ensure_space(total_h)
        for i, ln in enumerate(lines):
            c.setFillColor(colors.black)
            c.drawString(self.x0, self.y, f"- {ln}" if i == 0 else f"  {ln}")
            if i == 0:
                self._badge(status, self.x1, self.y)
            self.y -= leading * 0.85
        self.y -= 1 * mm

    def _photo_block(self, photos: list[bytes], keep_item_label: str | None = None):
        if not photos:
            return
        max_w = self.max_w
        max_h = 75 * mm
        gap = 5 * mm
        for b in photos:
            img = Image.open(io.BytesIO(b)).convert("RGB")
            iw, ih = img.size
            scale = min(max_w / iw, max_h / ih)
            tw, th = iw * scale, ih * scale
            needed = th + gap + 6 * mm
            if self.y - needed < self.content_bottom:
                self._new_page()
                if keep_item_label:
                    self._text_wrapped(f"(cont.) {keep_item_label}", size=9, bold=True, color=colors.grey, leading=11)
            x = self.x0
            y = self.y - th
            self.c.saveState()
            self.c.setStrokeColor(colors.lightgrey)
            self.c.setLineWidth(0.8)
            self.c.rect(x, y, tw, th, stroke=1, fill=0)
            img_buf = io.BytesIO()
            img.save(img_buf, format="JPEG", quality=85)
            img_buf.seek(0)
            self.c.drawImage(ImageReader(img_buf), x, y, width=tw, height=th, preserveAspectRatio=True, mask="auto")
            self.c.restoreState()
            self.y = y - gap

    def finalize(self):
        self._draw_footer()
        self.c.save()

def build_pdf_bytes(report: dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    pdf = PDFWriter(c, title="Trane Chiller Maintenance Report")
    header = report["header"]
    summary = compute_summary(report)

    pdf._text("Report Details", size=12, bold=True)
    pdf._ensure_space(20 * mm)
    rows = [
        ("Date", header.get("date", "")),
        ("Project", header.get("project", "")),
        ("Serial Number", header.get("serial_number", "")),
        ("Model", header.get("model", "")),
        ("Technician", header.get("technician", "")),
    ]
    c.setFont("Helvetica", 10)
    label_w = 28 * mm
    row_h = 6.5 * mm
    for k, v in rows:
        pdf._ensure_space(row_h)
        c.setFillColor(colors.grey)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(pdf.x0, pdf.y, f"{k}:")
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        c.drawString(pdf.x0 + label_w, pdf.y, safe_text(v) or "—")
        pdf.y -= row_h * 0.8
    pdf.y -= 4 * mm

    # Service Summary box
    box_h = 22 * mm
    pdf._ensure_space(box_h + 6 * mm)
    box_y = pdf.y - box_h
    c.saveState()
    c.setFillColor(colors.HexColor("#F7F7F7"))
    c.setStrokeColor(colors.lightgrey)
    c.setLineWidth(1)
    c.roundRect(pdf.x0, box_y, pdf.max_w, box_h, radius=3 * mm, stroke=1, fill=1)
    c.setFillColor(colors.black)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(pdf.x0 + 3 * mm, box_y + box_h - 6 * mm, "Service Summary")
    c.setFont("Helvetica-Bold", 10)
    c.drawRightString(pdf.x1 - 3 * mm, box_y + box_h - 6 * mm, f"Overall: {summary['overall']}")
    counts = summary["counts"]
    c.setFont("Helvetica", 9)
    c.drawString(pdf.x0 + 3 * mm, box_y + box_h - 12 * mm, f"OK: {counts['OK']} Not OK: {counts['Not OK']} N/A: {counts['N/A']}")
    not_ok_items = summary["not_ok_items"][:6]
    not_ok_text = ", ".join(not_ok_items) if not_ok_items else "None"
    c.setFillColor(colors.grey)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(pdf.x0 + 3 * mm, box_y + 5 * mm, "Not OK items:")
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 9)
    max_w = pdf.max_w - 22 * mm
    words = not_ok_text.split()
    line = []
    lines = []
    while words:
        line.append(words.pop(0))
        if words:
            w = c.stringWidth(" ".join(line + [words[0]]), "Helvetica", 9)
            if w > max_w:
                lines.append(" ".join(line))
                line = []
    if line:
        lines.append(" ".join(line))
    yy = box_y + 5 * mm
    for i, ln in enumerate(lines[:2]):
        c.drawString(pdf.x0 + 22 * mm, yy + (i * 4.2 * mm), ln)
    c.restoreState()
    pdf.y = box_y - 7 * mm

    pdf._text("Checklist", size=12, bold=True)
    pdf.y -= 2 * mm
    for sec in report["sections"]:
        pdf._section_bar(sec["title"])
        for item in sec["items"]:
            item_id = item["id"]
            label = item["label"]
            status = report["results"].get(item_id, {}).get("status", "") or "—"
            notes = safe_text(report["results"].get(item_id, {}).get("notes", ""))
            photos = (report.get("photos_by_item") or {}).get(item_id, [])
            pdf._item_line(label, status)
            if notes:
                pdf._text_wrapped(f"Notes: {notes}", size=9, color=colors.HexColor("#333333"), indent=8 * mm, leading=11)
            if photos:
                pdf._photo_block(photos, keep_item_label=label)
        pdf.y -= 1 * mm

    def write_bullets(title: str, text: str):
        pdf._section_bar(title)
        lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
        if not lines:
            pdf._text("• None", size=10)
            return
        for ln in lines:
            pdf._text_wrapped(f"• {ln}", size=10, leading=12)

    write_bullets("Findings / Issues", report.get("findings", ""))
    write_bullets("Recommendations / Actions", report.get("recommendations", ""))
    pdf.finalize()
    return buf.getvalue()

# DOCX Export
def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def build_docx_bytes(report: dict) -> bytes:
    doc = Document()
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = FOOTER_TEXT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)

    header = report["header"]
    summary = compute_summary(report)

    title = doc.add_paragraph("Trane Chiller Maintenance Report")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    doc.add_paragraph()

    h = doc.add_paragraph("Report Details")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)

    details = [
        ("Date", header.get("date", "") or "—"),
        ("Project", header.get("project", "") or "—"),
        ("Serial Number", header.get("serial_number", "") or "—"),
        ("Model", header.get("model", "") or "—"),
        ("Technician", header.get("technician", "") or "—"),
    ]
    t = doc.add_table(rows=len(details), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(details):
        t.cell(i, 0).text = f"{k}"
        t.cell(i, 1).text = f"{v}"
        t.cell(i, 0).paragraphs[0].runs[0].bold = True
        _set_cell_shading(t.cell(i, 0), "F2F2F2")

    doc.add_paragraph()
    sh = doc.add_paragraph("Service Summary")
    sh.runs[0].bold = True
    sh.runs[0].font.size = Pt(12)

    counts = summary["counts"]
    not_ok_list = summary["not_ok_items"] or ["None"]
    stbl = doc.add_table(rows=3, cols=2)
    stbl.style = "Table Grid"
    stbl.cell(0, 0).text = "Overall"
    stbl.cell(0, 1).text = summary["overall"]
    stbl.cell(1, 0).text = "Counts"
    stbl.cell(1, 1).text = f"OK: {counts['OK']} | Not OK: {counts['Not OK']} | N/A: {counts['N/A']}"
    stbl.cell(2, 0).text = "Not OK items"
    stbl.cell(2, 1).text = ", ".join(not_ok_list)
    for r in range(3):
        stbl.cell(r, 0).paragraphs[0].runs[0].bold = True
        _set_cell_shading(stbl.cell(r, 0), "F7F7F7")

    doc.add_paragraph()
    ch = doc.add_paragraph("Checklist")
    ch.runs[0].bold = True
    ch.runs[0].font.size = Pt(12)

    for sec in report["sections"]:
        secp = doc.add_paragraph(sec["title"])
        secp.runs[0].bold = True
        secp.runs[0].font.size = Pt(11)
        for item in sec["items"]:
            item_id = item["id"]
            label = item["label"]
            status = report["results"].get(item_id, {}).get("status", "") or "—"
            notes = safe_text(report["results"].get(item_id, {}).get("notes", ""))
            photos = (report.get("photos_by_item") or {}).get(item_id, [])
            p = doc.add_paragraph()
            p.add_run(f"- {label} ")
            rs = p.add_run(f"[{status}]")
            rs.bold = True
            if notes:
                pn = doc.add_paragraph(f"Notes: {notes}")
                pn.paragraph_format.left_indent = Inches(0.25)
                if pn.runs:
                    pn.runs[0].font.size = Pt(10)
            if photos:
                for b in photos:
                    img = Image.open(io.BytesIO(b)).convert("RGB")
                    img_buf = io.BytesIO()
                    img.save(img_buf, format="JPEG", quality=85)
                    img_buf.seek(0)
                    pic_par = doc.add_paragraph()
                    pic_par.paragraph_format.left_indent = Inches(0.25)
                    pic_par.add_run().add_picture(img_buf, width=Inches(6))
        doc.add_paragraph()

    def add_bullets(title: str, text: str):
        hh = doc.add_paragraph(title)
        hh.runs[0].bold = True
        hh.runs[0].font.size = Pt(12)
        lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
        if not lines:
            doc.add_paragraph("• None")
            return
        for ln in lines:
            doc.add_paragraph(f"• {ln}")

    add_bullets("Findings / Issues", report.get("findings", ""))
    add_bullets("Recommendations / Actions", report.get("recommendations", ""))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# XLSX Export
def build_xlsx_bytes(report: dict) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    header = report["header"]
    sections = report["sections"]
    ws.append(["Trane Chiller Maintenance Report"])
    ws.append([])
    ws.append(["Date", header["date"]])
    ws.append(["Project", header["project"]])
    ws.append(["Serial Number", header["serial_number"]])
    ws.append(["Model", header["model"]])
    ws.append(["Technician", header["technician"]])
    ws.append([])
    summary = compute_summary(report)
    ws.append(["Service Summary"])
    ws.append(["Overall", summary["overall"]])
    c = summary["counts"]
    ws.append(["Counts", f"OK: {c['OK']} | Not OK: {c['Not OK']} | N/A: {c['N/A']}"])
    ws.append(["Not OK items", ", ".join(summary["not_ok_items"]) or "None"])
    ws.append([])
    ws.append(["Checklist"])
    ws.append(["Section", "Item", "Status", "Notes"])
    for sec in sections:
        for item in sec["items"]:
            item_id = item["id"]
            label = item["label"]
            status = report["results"].get(item_id, {}).get("status", "")
            notes = report["results"].get(item_id, {}).get("notes", "")
            ws.append([sec["title"], label, status, notes])
    ws.append([])
    ws.append(["Findings / Issues", report["findings"].strip() or "None"])
    ws.append(["Recommendations / Actions", report["recommendations"].strip() or "None"])
    ws.append([])
    ws.append([FOOTER_TEXT])
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

# ────────────────────────────────────────────────
# Streamlit UI
# ────────────────────────────────────────────────
st.title("Trane Chiller Maintenance Report")
checklist_raw, checklist_msg = load_checklist(CHECKLIST_PATH)
sections = normalize_sections(checklist_raw)
st.caption(checklist_msg)

tab1, tab2 = st.tabs(["Create report", "Preview & export"])

if "results" not in st.session_state:
    st.session_state.results = {}
if "photos_by_item" not in st.session_state:
    st.session_state.photos_by_item = {}
if "current_report" not in st.session_state:
    st.session_state.current_report = None

with tab1:
    st.subheader("Report details")
    c1, c2 = st.columns(2)
    with c1:
        report_date = st.date_input("Date", value=date.today())
        project = st.text_input("Project")
        serial_number = st.text_input("Serial Number")
    with c2:
        model = st.text_input("Model")
        technician = st.text_input("Technician")

    st.divider()
    st.subheader("Checklist")

    for sec in sections:
        with st.expander(sec["title"], expanded=True):
            for item in sec["items"]:
                item_id = item["id"]
                label = item["label"]

                # Initialize
                if item_id not in st.session_state.results:
                    st.session_state.results[item_id] = {"status": "", "notes": ""}
                if item_id not in st.session_state.photos_by_item:
                    st.session_state.photos_by_item[item_id] = []

                colA, colB = st.columns([2, 3])
                with colA:
                    status = st.selectbox(
                        label,
                        options=STATUS_OPTIONS,
                        index=STATUS_OPTIONS.index(st.session_state.results[item_id]["status"])
                        if st.session_state.results[item_id]["status"] in STATUS_OPTIONS else 0,
                        key=f"status_{item_id}",
                    )
                with colB:
                    notes_label = "Notes (required for Not OK)" if status == "Not OK" else "Notes"
                    notes = st.text_input(
                        notes_label,
                        value=st.session_state.results[item_id]["notes"],
                        key=f"notes_{item_id}",
                        placeholder="Describe issue (required if Not OK)" if status == "Not OK" else "Optional notes",
                    )
                    if status == "Not OK" and not safe_text(notes):
                        st.warning("Notes are required when status is Not OK.", icon="⚠️")

                st.session_state.results[item_id]["status"] = status
                st.session_state.results[item_id]["notes"] = notes

                # ─── Photos ───────────────────────────────────────────────────────
                with st.expander("Photos for this item", expanded=False):
                    uploader_key = f"uploader_{item_id}"
                    state_key = f"uploader_state_{item_id}"

                    uploaded_files = st.file_uploader(
                        "Upload photos (jpg/png)",
                        type=["jpg", "jpeg", "png"],
                        accept_multiple_files=True,
                        key=uploader_key,
                    )

                    # Only process newly added files
                    if uploaded_files is not None:
                        if state_key not in st.session_state:
                            st.session_state[state_key] = []

                        prev_len = len(st.session_state[state_key])
                        current_len = len(uploaded_files)

                        if current_len > prev_len:
                            # New files were added
                            new_files = uploaded_files[prev_len:]
                            for file in new_files:
                                st.session_state.photos_by_item[item_id].append(file.getvalue())

                        # Update remembered state
                        st.session_state[state_key] = list(uploaded_files)  # copy reference

                    # Display current photos
                    current_photos = st.session_state.photos_by_item[item_id]
                    if current_photos:
                        cols = st.columns(3)
                        for idx, img_bytes in enumerate(current_photos):
                            with cols[idx % 3]:
                                st.image(img_bytes, use_container_width=True)

                    # Clear button
                    if current_photos and st.button("Clear all photos for this item", key=f"clear_{item_id}"):
                        st.session_state.photos_by_item[item_id] = []
                        if state_key in st.session_state:
                            del st.session_state[state_key]
                        st.rerun()

                st.divider()

    st.subheader("Findings / Issues")
    findings = st.text_area("Write each point on a new line", height=120, key="findings")

    st.subheader("Recommendations / Actions")
    recommendations = st.text_area("Write each point on a new line", height=120, key="recommendations")

    header = {
        "date": str(report_date),
        "project": safe_text(project),
        "serial_number": safe_text(serial_number),
        "model": safe_text(model),
        "technician": safe_text(technician),
    }

    st.session_state.current_report = {
        "header": header,
        "sections": sections,
        "results": st.session_state.results,
        "findings": findings,
        "recommendations": recommendations,
        "photos_by_item": st.session_state.photos_by_item,
    }

with tab2:
    report = st.session_state.get("current_report")
    if not report:
        st.info("Fill in the report first, then come back here.")
        st.stop()

    errors = validate_report(report)
    if errors:
        st.error("Fix the issues below before exporting:", icon="🛑")
        for e in errors:
            st.write(f"- {e}")

    st.subheader("Preview")
    h = report["header"]
    summary = compute_summary(report)
    cA, cB, cC = st.columns(3)
    with cA:
        st.markdown(f"**Date:** {h['date'] or '—'}")
        st.markdown(f"**Project:** {h['project'] or '—'}")
    with cB:
        st.markdown(f"**Serial Number:** {h['serial_number'] or '—'}")
        st.markdown(f"**Model:** {h['model'] or '—'}")
    with cC:
        st.markdown(f"**Technician:** {h['technician'] or '—'}")
        st.markdown(f"**Overall:** {summary['overall']}")
    st.caption(f"Counts — OK: {summary['counts']['OK']}, Not OK: {summary['counts']['Not OK']}, N/A: {summary['counts']['N/A']}")
    st.divider()

    for sec in report["sections"]:
        st.markdown(f"### {sec['title']}")
        for item in sec["items"]:
            item_id = item["id"]
            status = report["results"].get(item_id, {}).get("status", "") or "—"
            notes = safe_text(report["results"].get(item_id, {}).get("notes", ""))
            st.write(f"- {item['label']} — **{status}**")
            if notes:
                st.caption(f"Notes: {notes}")
            photos = (report.get("photos_by_item") or {}).get(item_id, [])
            if photos:
                cols = st.columns(3)
                for i, b in enumerate(photos):
                    with cols[i % 3]:
                        st.image(b, use_container_width=True)

    st.markdown("### Findings / Issues")
    lines = [ln.strip() for ln in (report.get("findings") or "").splitlines() if ln.strip()]
    if not lines:
        st.write("• None")
    else:
        for ln in lines:
            st.write(f"• {ln}")

    st.markdown("### Recommendations / Actions")
    lines = [ln.strip() for ln in (report.get("recommendations") or "").splitlines() if ln.strip()]
    if not lines:
        st.write("• None")
    else:
        for ln in lines:
            st.write(f"• {ln}")

    st.divider()
    st.subheader("Export")
    file_base = f"trane_chiller_report_{h['date'] or 'report'}"
    disabled = bool(errors)
    if disabled:
        st.info("Export buttons are disabled until all validation issues are fixed.")

    pdf_bytes = build_pdf_bytes(report)
    docx_bytes = build_docx_bytes(report)
    xlsx_bytes = build_xlsx_bytes(report)

    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name=f"{file_base}.pdf",
            mime="application/pdf",
            disabled=disabled,
        )
    with c2:
        st.download_button(
            "Download Word (DOCX)",
            data=docx_bytes,
            file_name=f"{file_base}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=disabled,
        )
    with c3:
        st.download_button(
            "Download Excel",
            data=xlsx_bytes,
            file_name=f"{file_base}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            disabled=disabled,
        )
